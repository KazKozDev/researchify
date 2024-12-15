from typing import Dict, Optional, Union, List
import PyPDF2
import docx
import pandas as pd
import openpyxl
import chardet
import io
import os
from datetime import datetime
from werkzeug.utils import secure_filename
import logging
from vector_store import Document, VectorStore, RAGSystem
from Crypto.Cipher import AES

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handle document processing, display, and storage for various file formats."""
    
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'word': ['.docx', '.doc'],
        'excel': ['.xlsx', '.xls', '.csv'],
        'text': ['.txt', '.md', '.rst']
    }
    
    def __init__(
        self,
        vector_store: VectorStore,
        rag_system: RAGSystem,
        upload_folder: str = 'uploads',
        max_file_size: int = 10 * 1024 * 1024  # 10MB
    ):
        """
        Initialize the document processor.
        
        Args:
            vector_store: VectorStore instance for document storage
            rag_system: RAGSystem instance for document retrieval
            upload_folder: Directory for temporary file storage
            max_file_size: Maximum allowed file size in bytes
        """
        self.vector_store = vector_store
        self.rag_system = rag_system
        self.upload_folder = upload_folder
        self.max_file_size = max_file_size
        
        # Create upload directory if it doesn't exist
        os.makedirs(upload_folder, exist_ok=True)
        
    def _check_file_size(self, file) -> bool:
        """Check if file size is within limits."""
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        return size <= self.max_file_size
    
    def _get_file_format(self, filename: str) -> Optional[str]:
        """Determine file format from extension."""
        ext = os.path.splitext(filename)[1].lower()
        for format_type, extensions in self.SUPPORTED_FORMATS.items():
            if ext in extensions:
                return format_type
        return None
    
    def _read_pdf(self, file) -> Dict:
        """Extract text and metadata from PDF files."""
        try:
            pdf_reader = PyPDF2.PdfReader(file, strict=False)
            
            # Handle encrypted PDFs
            if pdf_reader.is_encrypted:
                try:
                    pdf_reader.decrypt('')  # Try empty password first
                except:
                    raise ValueError("PDF is encrypted and requires a password")
            
            text = ""
            for page in pdf_reader.pages:
                try:
                    text += page.extract_text() + "\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page: {str(e)}")
                    continue
                
            metadata = pdf_reader.metadata or {}
            return {
                'content': text,
                'metadata': {
                    'title': metadata.get('/Title', ''),
                    'author': metadata.get('/Author', ''),
                    'creation_date': metadata.get('/CreationDate', ''),
                    'pages': len(pdf_reader.pages),
                    'is_encrypted': pdf_reader.is_encrypted
                }
            }
        except Exception as e:
            logger.error(f"Error reading PDF: {str(e)}")
            raise
    
    def _read_word(self, file) -> Dict:
        """Extract text and metadata from Word documents."""
        try:
            doc = docx.Document(file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            core_properties = doc.core_properties
            return {
                'content': text,
                'metadata': {
                    'title': core_properties.title or '',
                    'author': core_properties.author or '',
                    'created': str(core_properties.created or ''),
                    'modified': str(core_properties.modified or ''),
                    'paragraphs': len(doc.paragraphs)
                }
            }
        except Exception as e:
            logger.error(f"Error reading Word document: {str(e)}")
            raise
    
    def _read_excel(self, file) -> Dict:
        """Extract data from Excel files and CSV."""
        try:
            try:
                # Try reading as Excel first
                df = pd.read_excel(file)
            except:
                # If fails, try as CSV
                file.seek(0)
                df = pd.read_csv(file)
                
            # Convert DataFrame to formatted string
            text = (
                f"Columns: {', '.join(df.columns)}\n\n"
                f"Data Preview:\n{df.head().to_string()}\n\n"
                f"Summary Statistics:\n{df.describe().to_string()}"
            )
            
            return {
                'content': text,
                'metadata': {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': list(df.columns),
                    'file_type': 'csv' if file.filename.endswith('.csv') else 'excel'
                }
            }
        except Exception as e:
            logger.error(f"Error reading Excel/CSV file: {str(e)}")
            raise
    
    def _read_text(self, file) -> Dict:
        """Read plain text files with encoding detection."""
        try:
            content = file.read()
            if isinstance(content, bytes):
                encoding_info = chardet.detect(content)
                encoding = encoding_info['encoding'] or 'utf-8'
                content = content.decode(encoding)
                
            return {
                'content': content,
                'metadata': {
                    'encoding': encoding if isinstance(content, str) else 'utf-8',
                    'length': len(content),
                    'lines': content.count('\n') + 1
                }
            }
        except Exception as e:
            logger.error(f"Error reading text file: {str(e)}")
            raise
    
    def process_file(self, file) -> Dict:
        """
        Process uploaded file and return its content and display format.
        
        Args:
            file: File object from request.files
            
        Returns:
            Dictionary containing processed content and display information
        """
        if not file or not file.filename:
            raise ValueError("No file provided")
            
        if not self._check_file_size(file):
            raise ValueError(f"File size exceeds maximum limit of {self.max_file_size/1024/1024}MB")
        
        filename = secure_filename(file.filename)
        file_format = self._get_file_format(filename)
        
        if not file_format:
            raise ValueError(f"Unsupported file format: {filename}")
        
        # Process file based on format
        try:
            logger.info(f"Processing file: {filename} (format: {file_format})")
            
            if file_format == 'pdf':
                result = self._read_pdf(file)
            elif file_format == 'word':
                result = self._read_word(file)
            elif file_format == 'excel':
                result = self._read_excel(file)
            else:  # text
                result = self._read_text(file)
                
            # Save file
            file_path = os.path.join(self.upload_folder, filename)
            file.seek(0)
            file.save(file_path)
            
            # Add timestamp to metadata
            result['metadata']['processed_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            # Generate display HTML
            display_html = self._generate_display_html(result, file_format, filename)
            
            # Add to RAG system
            self._add_to_rag(result, filename)
            
            logger.info(f"Successfully processed file: {filename}")
            
            return {
                'success': True,
                'display_html': display_html,
                'format': file_format,
                'filename': filename,
                'metadata': result['metadata']
            }
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            raise
    
    def _generate_display_html(self, result: Dict, file_format: str, filename: str) -> str:
        """Generate HTML for displaying file content in chat."""
        content_preview = result['content'][:1000] + "..." if len(result['content']) > 1000 else result['content']
        
        html = f"""
        <div class="document-preview p-4 border rounded-lg mb-4">
            <div class="document-header flex justify-between items-center mb-4">
                <h3 class="text-lg font-medium">{filename}</h3>
                <span class="text-sm text-gray-500">{file_format.upper()}</span>
            </div>
            <div class="metadata mb-4 text-sm text-gray-600">
                <h4 class="font-medium mb-2">Document Information:</h4>
                <ul class="list-disc pl-4">
        """
        
        for key, value in result['metadata'].items():
            html += f"<li><strong>{key.title()}:</strong> {value}</li>"
        
        html += "</ul></div>"
        
        # Add content preview based on format
        if file_format == 'excel':
            html += f"""
                <div class="content-preview">
                    <pre class="bg-gray-50 p-4 rounded overflow-x-auto">
                        {content_preview}
                    </pre>
                </div>
            """
        else:
            html += f"""
                <div class="content-preview whitespace-pre-wrap">
                    {content_preview}
                </div>
            """
        
        html += "</div>"
        return html
    
    def _add_to_rag(self, result: Dict, filename: str) -> None:
        """Add processed document to RAG system."""
        doc = Document(
            id=f"doc_{filename}_{datetime.now().timestamp()}",
            title=filename,
            authors=result['metadata'].get('author', ''),
            abstract=result['content'][:500],  # Use first 500 chars as abstract
            pdf_link=None,
            arxiv_link=None,
            published=result['metadata'].get('creation_date', ''),
            categories='uploaded_document'
        )
        
        self.vector_store.add_documents([doc])
        logger.info(f"Added document to RAG system: {filename}")
    
    def analyze_content(self, filename: str) -> str:
        """Analyze the content of a specific file."""
        file_path = os.path.join(self.upload_folder, filename)
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {filename}")
            
        # Get file format
        file_format = self._get_file_format(filename)
        if not file_format:
            raise ValueError(f"Unsupported file format: {filename}")
            
        # Read and analyze file
        try:
            with open(file_path, 'rb') as file:
                if file_format == 'pdf':
                    result = self._read_pdf(file)
                elif file_format == 'word':
                    result = self._read_word(file)
                elif file_format == 'excel':
                    result = self._read_excel(file)
                else:
                    result = self._read_text(file)
                    
            # Generate analysis
            analysis = (
                f"File Analysis for {filename}\n\n"
                f"Format: {file_format.upper()}\n"
                f"Content Length: {len(result['content'])} characters\n\n"
                "Metadata:\n"
            )
            
            for key, value in result['metadata'].items():
                analysis += f"- {key.title()}: {value}\n"
                
            logger.info(f"Successfully analyzed file: {filename}")
            return analysis
            
        except Exception as e:
            logger.error(f"Error analyzing file {filename}: {str(e)}")
            raise